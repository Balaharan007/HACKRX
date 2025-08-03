import os
import requests
import PyPDF2
import docx
from io import BytesIO
from typing import List, Dict, Any, Union
import pinecone
from pinecone import Pinecone
import google.generativeai as genai
from dotenv import load_dotenv
import hashlib
import json
import re
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import tempfile

load_dotenv()

class DocumentProcessor:
    def __init__(self):
        # Initialize Pinecone
        self.pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
        self.index_name = "hackrx-documents"
        
        # Create index if it doesn't exist
        try:
            self.index = self.pc.Index(self.index_name)
        except:
            try:
                self.pc.create_index(
                    name=self.index_name,
                    dimension=384,  # for TF-IDF vectors (we'll use 384 dimensions)
                    metric="cosine",
                    spec=pinecone.ServerlessSpec(
                        cloud="aws",
                        region="us-east-1"
                    )
                )
                self.index = self.pc.Index(self.index_name)
            except Exception as e:
                print(f"Pinecone setup error: {e}")
                self.index = None
        
        # Initialize TF-IDF vectorizer for embeddings
        self.vectorizer = TfidfVectorizer(
            max_features=384,
            stop_words='english',
            ngram_range=(1, 2),
            max_df=0.95,
            min_df=2
        )
        self.is_vectorizer_fitted = False
        
        # Initialize Gemini
        genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
        self.gemini_model = genai.GenerativeModel('gemini-2.0-flash-exp')
    
    def download_document(self, url: str) -> bytes:
        """Download document from URL"""
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            return response.content
        except Exception as e:
            raise Exception(f"Failed to download document: {str(e)}")
    
    def extract_text_from_pdf(self, content: Union[bytes, str]) -> str:
        """Extract text from PDF"""
        try:
            if isinstance(content, str):
                # If it's a file path
                with open(content, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
            else:
                # If it's bytes content
                pdf_file = BytesIO(content)
                reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise Exception(f"Failed to extract PDF text: {str(e)}")
    
    def extract_text_from_docx(self, content: Union[bytes, str]) -> str:
        """Extract text from DOCX"""
        try:
            if isinstance(content, str):
                # If it's a file path
                doc = docx.Document(content)
            else:
                # If it's bytes content
                doc = docx.Document(BytesIO(content))
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Failed to extract DOCX text: {str(e)}")
    
    def extract_text_from_txt(self, content: Union[bytes, str]) -> str:
        """Extract text from TXT file"""
        try:
            if isinstance(content, str):
                # If it's a file path
                with open(content, 'r', encoding='utf-8') as file:
                    return file.read()
            else:
                # If it's bytes content
                return content.decode('utf-8')
        except Exception as e:
            raise Exception(f"Failed to extract TXT text: {str(e)}")
    
    def process_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """Process uploaded file content"""
        file_ext = filename.lower().split('.')[-1]
        
        if file_ext == 'pdf':
            return self.extract_text_from_pdf(file_content)
        elif file_ext in ['docx', 'doc']:
            return self.extract_text_from_docx(file_content)
        elif file_ext == 'txt':
            return self.extract_text_from_txt(file_content)
        else:
            # Try to decode as text
            try:
                return file_content.decode('utf-8')
            except:
                raise Exception(f"Unsupported file format: {file_ext}")
    
    def extract_text(self, source: str, is_file_path: bool = False) -> str:
        """Extract text from document URL or file path"""
        if is_file_path:
            # Handle local file
            if source.lower().endswith('.pdf'):
                return self.extract_text_from_pdf(source)
            elif source.lower().endswith(('.docx', '.doc')):
                return self.extract_text_from_docx(source)
            elif source.lower().endswith('.txt'):
                return self.extract_text_from_txt(source)
            else:
                raise Exception("Unsupported file format")
        else:
            # Handle URL
            content = self.download_document(source)
            
            if source.lower().endswith('.pdf'):
                return self.extract_text_from_pdf(content)
            elif source.lower().endswith(('.docx', '.doc')):
                return self.extract_text_from_docx(content)
            elif source.lower().endswith('.txt'):
                return self.extract_text_from_txt(content)
            else:
                # Try to decode as plain text
                try:
                    return content.decode('utf-8')
                except:
                    raise Exception("Unsupported document format")
    
    def chunk_text(self, text: str, chunk_size: int = 1500, overlap: int = 300) -> List[Dict]:
        """Split text into chunks with metadata"""
        # Clean text
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Split into sentences first
        sentences = re.split(r'[.!?]+', text)
        
        chunks = []
        current_chunk = ""
        chunk_id = 0
        start_pos = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # Check if adding this sentence would exceed chunk size
            if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                # Save current chunk
                chunks.append({
                    "id": f"chunk_{chunk_id}",
                    "text": current_chunk.strip(),
                    "start_pos": start_pos,
                    "end_pos": start_pos + len(current_chunk)
                })
                
                # Start new chunk with overlap
                overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                current_chunk = overlap_text + " " + sentence
                start_pos = start_pos + len(current_chunk) - len(overlap_text) - len(sentence)
                chunk_id += 1
            else:
                current_chunk += " " + sentence if current_chunk else sentence
        
        # Add the last chunk
        if current_chunk.strip():
            chunks.append({
                "id": f"chunk_{chunk_id}",
                "text": current_chunk.strip(),
                "start_pos": start_pos,
                "end_pos": start_pos + len(current_chunk)
            })
        
        return chunks
    
    def create_embeddings(self, texts: List[str]) -> np.ndarray:
        """Create embeddings using TF-IDF"""
        if not texts:
            return np.array([]).reshape(0, 384)
        
        try:
            if not self.is_vectorizer_fitted:
                # Fit the vectorizer on the texts
                embeddings = self.vectorizer.fit_transform(texts)
                self.is_vectorizer_fitted = True
            else:
                # Transform using already fitted vectorizer
                embeddings = self.vectorizer.transform(texts)
            
            # Convert to dense array and pad/truncate to 384 dimensions
            dense_embeddings = embeddings.toarray()
            
            # Ensure 384 dimensions
            if dense_embeddings.shape[1] < 384:
                # Pad with zeros
                padding = np.zeros((dense_embeddings.shape[0], 384 - dense_embeddings.shape[1]))
                dense_embeddings = np.hstack([dense_embeddings, padding])
            elif dense_embeddings.shape[1] > 384:
                # Truncate
                dense_embeddings = dense_embeddings[:, :384]
            
            return dense_embeddings
            
        except Exception as e:
            print(f"Embedding creation error: {e}")
            # Fallback: create random embeddings
            return np.random.rand(len(texts), 384)
    
    def embed_chunks(self, chunks: List[Dict]) -> List[Dict]:
        """Generate embeddings for chunks"""
        texts = [chunk["text"] for chunk in chunks]
        embeddings = self.create_embeddings(texts)
        
        for i, chunk in enumerate(chunks):
            if i < len(embeddings):
                chunk["embedding"] = embeddings[i].tolist()
            else:
                chunk["embedding"] = np.random.rand(384).tolist()
        
        return chunks
    
    def store_in_pinecone(self, chunks: List[Dict], document_id: str):
        """Store chunks in Pinecone"""
        if not self.index:
            print("Pinecone index not available, skipping storage")
            return
            
        vectors = []
        for chunk in chunks:
            vector_id = f"{document_id}_{chunk['id']}"
            # Limit metadata text size
            metadata_text = chunk["text"]
            if len(metadata_text) > 900:  # Pinecone metadata limit
                metadata_text = metadata_text[:900] + "..."
                
            vectors.append({
                "id": vector_id,
                "values": chunk["embedding"],
                "metadata": {
                    "document_id": document_id,
                    "chunk_id": chunk["id"],
                    "text": metadata_text,
                    "start_pos": chunk["start_pos"],
                    "end_pos": chunk["end_pos"]
                }
            })
        
        # Upsert in batches
        batch_size = 100
        for i in range(0, len(vectors), batch_size):
            batch = vectors[i:i + batch_size]
            try:
                self.index.upsert(vectors=batch)
                print(f"Stored batch {i//batch_size + 1} of {len(vectors)//batch_size + 1}")
            except Exception as e:
                print(f"Pinecone upsert error: {e}")
    
    def process_document(self, source: str, is_file_path: bool = False, file_content: bytes = None, filename: str = None) -> Dict:
        """Complete document processing pipeline"""
        try:
            # Extract text
            if file_content and filename:
                # Process uploaded file
                text = self.process_uploaded_file(file_content, filename)
                document_id = hashlib.md5((filename + str(len(file_content))).encode()).hexdigest()
            else:
                # Process URL or file path
                text = self.extract_text(source, is_file_path)
                document_id = hashlib.md5(source.encode()).hexdigest()
            
            # Chunk text
            chunks = self.chunk_text(text)
            
            # Generate embeddings
            chunks = self.embed_chunks(chunks)
            
            # Store in Pinecone
            self.store_in_pinecone(chunks, document_id)
            
            return {
                "success": True,
                "text": text,
                "chunks": len(chunks),
                "document_id": document_id,
                "message": f"Successfully processed document with {len(chunks)} chunks"
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def search_similar_chunks(self, query: str, document_id: str = None, top_k: int = 5) -> List[Dict]:
        """Search for similar chunks"""
        try:
            # Create query embedding
            query_embedding = self.create_embeddings([query])[0].tolist()
            
            filter_dict = {}
            if document_id:
                filter_dict = {"document_id": document_id}
            
            if self.index:
                search_response = self.index.query(
                    vector=query_embedding,
                    top_k=top_k,
                    include_metadata=True,
                    filter=filter_dict if filter_dict else None
                )
                
                results = []
                for match in search_response.matches:
                    results.append({
                        "score": match.score,
                        "text": match.metadata.get("text", ""),
                        "chunk_id": match.metadata.get("chunk_id", "unknown"),
                        "document_id": match.metadata.get("document_id", "unknown")
                    })
                
                return results
            else:
                # Fallback: return empty results
                return []
        
        except Exception as e:
            print(f"Search error: {e}")
            return []
    
    def generate_answer(self, question: str, relevant_chunks: List[Dict]) -> Dict:
        """Generate answer using Gemini"""
        if not relevant_chunks:
            return {
                "answer": "I couldn't find relevant information in the document to answer your question.",
                "relevant_chunks": [],
                "reasoning": "No relevant document sections found"
            }
            
        context = "\n\n".join([f"[Clause {chunk['chunk_id']}]: {chunk['text']}" for chunk in relevant_chunks])
        
        prompt = f"""
        You are an intelligent document analysis agent specializing in insurance, legal, HR, and compliance domains.
        
        Based on the following document clauses and user question, provide a comprehensive and accurate answer.

        DOCUMENT CLAUSES:
        {context}

        USER QUESTION:
        {question}

        INSTRUCTIONS:
        1. Analyze the relevant clauses carefully
        2. Provide a clear, accurate answer based ONLY on the document content
        3. If the document doesn't contain sufficient information, state that clearly
        4. Reference specific clauses when relevant
        5. Be precise and professional
        6. Focus on insurance policies, legal contracts, HR policies, or compliance documents
        7. Provide explainable reasoning for your decision
        8. Answer should be concise but comprehensive
        9. Cite specific clause numbers or sections when possible

        ANSWER:
        """
        
        try:
            response = self.gemini_model.generate_content(prompt)
            return {
                "answer": response.text,
                "relevant_chunks": relevant_chunks,
                "reasoning": f"Answer based on semantic similarity search of {len(relevant_chunks)} document clauses"
            }
        except Exception as e:
            return {
                "answer": f"Error generating answer: {str(e)}",
                "relevant_chunks": relevant_chunks,
                "reasoning": "Error occurred during LLM processing"
            }
